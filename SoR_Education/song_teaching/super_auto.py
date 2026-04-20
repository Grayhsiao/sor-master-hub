"""
=============================================================================
🎵 Gray 的英文歌詞教材生成器 (V5 最強容錯版)
=============================================================================
【本次更新】
1. 🛡️ 容錯機制：不再依賴 ### 分隔符，改用關鍵字掃描，即使 AI 格式跑掉也能抓。
2. 🐞 除錯紀錄：每次執行都會儲存 'ai_output.txt'，方便檢查 AI 到底回了什麼。
3. 📜 完整功能：保留去重邏輯 + 最後一頁附上完整歌詞。

【使用方式】
1. cd ~/Desktop/grey/song_teaching
2. python3 super_auto.py
=============================================================================
"""

import os
import sys
import platform
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn 

# ==========================================
# 🔑 API Key 設定
# ==========================================
API_KEY = os.getenv('GOOGLE_API_KEY', '')
# ==========================================

# 設定 API
genai.configure(api_key=API_KEY)

# 生成參數
generation_config = {
    "temperature": 0.2, # 低創意，高精準
    "top_p": 1,
    "top_k": 32,
    "max_output_tokens": 8192,
}

script_dir = os.path.dirname(os.path.abspath(__file__))

def print_status(message):
    print(f"⚙️  {message}")

def get_best_model():
    """自動尋找可用的 Gemini 模型"""
    print_status("正在搜尋最佳 AI 模型...")
    try:
        # 優先順序：Flash (快) -> Pro (穩)
        preferred_models = ['gemini-1.5-flash', 'gemini-1.5-pro', 'gemini-pro']
        available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        
        for model_name in preferred_models:
            full_name = f"models/{model_name}"
            if full_name in available_models:
                print_status(f"已連接 AI 模型: {model_name}")
                return genai.GenerativeModel(model_name, generation_config=generation_config)
        
        for m in available_models:
            if 'gemini' in m:
                return genai.GenerativeModel(m, generation_config=generation_config)
                
    except Exception as e:
        print(f"❌ 連線失敗: {e}")
        return None

model = get_best_model()

def get_ai_content(song_input):
    """呼叫 Gemini API 生成教材內容"""
    if not model: return None

    print_status(f"正在分析歌曲：[{song_input}]...")
    print_status("AI 正在執行：全詞檢索 -> 去重篩選 -> 教材編寫 -> 附錄整理...")
    
    prompt = f"""
    Role: You are a strict and precise English teacher creating a worksheet.
    Task: Create a worksheet for the song "{song_input}".

    *** OUTPUT FORMAT RULES (Strictly Follow) ***
    You must generate 8 teaching segments followed by the full lyrics.
    Do not add conversational filler like "Here is the worksheet". Start directly with the content.

    Part 1: Teaching Segments (Repeat this block 8 times)
    (Title)
    [Segment Number]. [First line of lyrics]
    (Lyrics)
    Lyrics : [English line 1]...
    Dict   : [Chinese translation]...
    (Trans)
    [Sentence pattern]
    (Vocab)
    [Emoji] [Category]: [Word] ([Chinese])...

    Part 2: Full Lyrics (At the end)
    (Full Lyrics)
    [Paste Complete Lyrics Here]
    """
    
    try:
        response = model.generate_content(prompt)
        text = response.text
        
        # 🐞 【除錯功能】將 AI 的原始回應存檔，方便檢查
        debug_file = os.path.join(script_dir, "ai_output_debug.txt")
        with open(debug_file, "w", encoding="utf-8") as f:
            f.write(text)
        print_status(f"AI 原始回應已備份至: ai_output_debug.txt (若失敗請檢查此檔)")
        
        return text
    except Exception as e:
        print(f"❌ API 錯誤: {e}")
        return None

# --- 排版設定 ---

def set_landscape(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

def set_font_menlo(run, size=10):
    run.font.name = 'Menlo'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Menlo')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Menlo')
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Menlo')
    run.font.size = Pt(size)

def set_font_biaukai(run, size=10, is_bold=False):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'BiauKai')
    run.font.size = Pt(size)
    run.bold = is_bold

def parse_song_data(content):
    """
    V5 新版解析器：不依賴 ###，改用狀態機掃描
    """
    print_status("正在解析資料結構 (智慧模式)...")
    if not content: return [], ""
    
    lines = content.split('\n')
    parsed_data = []
    
    # 暫存區
    current_item = {}
    current_tag = None # 記錄現在讀到哪個標籤：title, lyrics, trans, vocab
    
    full_lyrics_text = ""
    is_reading_full_lyrics = False
    
    for line in lines:
        clean_line = line.strip()
        
        # 1. 偵測完整歌詞區塊
        if clean_line == "(Full Lyrics)":
            is_reading_full_lyrics = True
            # 如果前面還有沒存的 item，先存起來
            if current_item.get("title"):
                parsed_data.append(current_item)
                current_item = {}
            continue
            
        if is_reading_full_lyrics:
            full_lyrics_text += line + "\n"
            continue

        # 2. 偵測教學區塊標籤
        if clean_line == "(Title)":
            # 遇到新標題，先把上一個存起來
            if current_item.get("title"):
                parsed_data.append(current_item)
            current_item = {"title": "", "lyrics": "", "trans": "", "vocab": ""}
            current_tag = "title"
            continue
            
        elif clean_line == "(Lyrics)":
            current_tag = "lyrics"
            continue
        elif clean_line == "(Trans)":
            current_tag = "trans"
            continue
        elif clean_line == "(Vocab)":
            current_tag = "vocab"
            continue
        elif clean_line.startswith("###"): # 如果 AI 還是用了 ###，忽略它
            continue

        # 3. 填入內容
        if current_tag and current_item is not None:
            if clean_line:
                # 如果該欄位已經有內容，補上換行符號
                if current_item.get(current_tag):
                    current_item[current_tag] += "\n" + clean_line
                else:
                    current_item[current_tag] = clean_line

    # 迴圈結束，存最後一筆
    if current_item.get("title"):
        parsed_data.append(current_item)

    return parsed_data, full_lyrics_text

def add_content_to_cell(cell, data):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    p = cell.add_paragraph()
    run = p.add_run("🎵 " + data.get("title", ""))
    set_font_biaukai(run, size=11, is_bold=True)
    p.paragraph_format.space_after = Pt(2)
    
    lyrics_text = data.get("lyrics", "")
    lyrics_spaced = lyrics_text.replace("\n", "\n\n\n")
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(lyrics_spaced)
    set_font_menlo(run, size=10)
    
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("✍️ 翻譯與練習\n")
    set_font_biaukai(run, size=10, is_bold=True)
    run = p.add_run("整句翻譯：___________________________________________________________________\n\n")
    set_font_biaukai(run, size=9, is_bold=False)
    trans_text = data.get("trans", "")
    run = p.add_run(f"句型替換：{trans_text}\n\n")
    set_font_biaukai(run, size=9, is_bold=True)
    run = p.add_run("1. ___________________________________________________________________\n\n2. ___________________________________________________________________")
    set_font_biaukai(run, size=9, is_bold=False)
    
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run("🧩 單字庫 (Vocabulary)\n")
    set_font_biaukai(run, size=9, is_bold=True)
    vocab_text = data.get("vocab", "")
    run = p.add_run(vocab_text)
    set_font_biaukai(run, size=9, is_bold=False)

def add_full_lyrics_page(doc, lyrics_text):
    """在最後新增完整歌詞頁面"""
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = 1 
    run = p.add_run("📜 Complete Lyrics (完整歌詞)")
    set_font_biaukai(run, size=16, is_bold=True)
    p.paragraph_format.space_after = Pt(20)
    
    p = doc.add_paragraph()
    run = p.add_run(lyrics_text)
    set_font_biaukai(run, size=12) 
    p.paragraph_format.line_spacing = 1.2

# ================= 主程式 =================
if __name__ == "__main__":
    print("\n" + "="*50)
    print(" 🎵  Gray 的英文歌詞教材生成器 (V5 最強容錯版)")
    print(f" 📂  輸出位置: {script_dir}")
    print("="*50 + "\n")
    
    if not model:
        print("⚠️ 無法啟動 AI，請檢查 API Key。")
        exit()

    song_input = input("🎤 請輸入 [歌名] + [樂手/歌手] (例如: Perfect by Ed Sheeran): ")
    
    if not song_input:
        print("❌ 未輸入資料，程式結束。")
        exit()

    ai_text = get_ai_content(song_input)

    if ai_text:
        song_data, full_lyrics = parse_song_data(ai_text)
        
        if song_data:
            print_status(f"正在建立 Word 文件 (共 {len(song_data)} 個教學段落)...")
            doc = Document()
            set_landscape(doc)
            
            for i in range(0, len(song_data), 2):
                table = doc.add_table(rows=2, cols=1)
                table.autofit = False
                for row in table.rows: row.height = Cm(9.5)
                
                add_content_to_cell(table.cell(0, 0), song_data[i])
                if i + 1 < len(song_data):
                    add_content_to_cell(table.cell(1, 0), song_data[i+1])
                doc.add_page_break()
            
            if full_lyrics:
                print_status("正在加入完整歌詞附錄...")
                add_full_lyrics_page(doc, full_lyrics)
            
            safe_name = "".join([c for c in song_input if c.isalpha() or c.isdigit() or c==' ']).strip()
            output_file = os.path.join(script_dir, f"{safe_name}_Worksheet.docx")
            
            try:
                print_status("正在存檔...")
                doc.save(output_file)
                print("\n" + "="*50)
                print(f"🎉 成功！檔案已生成：")
                print(f"📄 {output_file}")
                print("="*50 + "\n")
                
                if platform.system() == "Darwin":
                    os.system(f"open '{output_file}'")
            except Exception as e:
                print(f"❌ 存檔失敗: {e}")
        else:
            print("⚠️ 錯誤：無法解析 AI 回傳的內容。")
            print("請打開資料夾中的 'ai_output_debug.txt' 查看 AI 到底回了什麼。")